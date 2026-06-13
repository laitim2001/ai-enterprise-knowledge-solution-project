"""W72-prep 勘查:抽 miss 文件真內容驗 ground truth + PDF text-layer 信號探測。

用 pypdf(唔 OCR)抽 PDF text-layer → 探測 image-based vs text-based(額外 PDF 信號);
用 python-docx 抽 docx 頭幾段 → 確認 form / prose / 簡報 嘅人類判斷啱唔啱。
全部秒級(繞開 v1 嗰個 35 分鐘 OCR 成本)。
"""

from __future__ import annotations

from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SAMPLE = ROOT / "docs" / "06-reference" / "01-sample-doc"


def read_pdf(path: Path, pages: int = 4):
    try:
        import pypdfium2 as pdfium
    except Exception as e:  # noqa: BLE001
        return f"[pypdfium2 不可用: {e}]"
    try:
        pdf = pdfium.PdfDocument(str(path))
        n = len(pdf)
        lines = [f"  總頁數={n}"]
        empty = 0
        for i in range(min(pages, n)):
            page = pdf[i]
            tp = page.get_textpage()
            t = (tp.get_text_range() or "").strip()
            if not t:
                empty += 1
            lines.append(f"  --- page {i + 1} text-layer chars={len(t)} ---")
            lines.append("    " + (t[:400].replace("\n", " ") if t else "[空 → image-based page]"))
        lines.append(f"  >>> 頭 {min(pages, n)} 頁有 {empty} 頁 text-layer 空(空多=image/scan-based)")
        return "\n".join(lines)
    except Exception as e:  # noqa: BLE001
        return f"  [pypdfium2 parse err: {e}]"


def read_docx(path: Path, n: int = 18):
    try:
        import docx
    except Exception as e:  # noqa: BLE001
        return f"[python-docx 不可用: {e}]"
    try:
        d = docx.Document(str(path))
        paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        tbls = len(d.tables)
        lines = [f"  非空段落={len(paras)}  tables={tbls}"]
        for i, t in enumerate(paras[:n]):
            lines.append(f"  {i:2}: {t[:110]}")
        return "\n".join(lines)
    except Exception as e:  # noqa: BLE001
        return f"  [docx err: {e}]"


TARGETS = [
    ("AI demonstration_20260329 (4).pdf", "pdf", "人判 P3 簡報?(rule 撞 P1)"),
    ("AI-Governance-Procedure.pdf", "pdf", "人判 P1 procedure(對照)"),
    ("Virtual_Boss_3.5_Information_Intake_Questionnaire.docx", "docx", "人判 FORM 問卷?(rule 撞 P2)"),
    ("Virtual_Boss_3.5_Introduction_and_Brief.docx", "docx", "人判 P2 短文?(rule 撞 too_small)"),
    ("AI_Project_Registration_Form.docx", "docx", "人判 FORM(對照,已命中)"),
]


def main():
    for name, fmt, note in TARGETS:
        path = SAMPLE / name
        print(f"\n{'=' * 70}\n[{fmt}] {name}\n  ({note})\n{'=' * 70}")
        if not path.exists():
            print("  [檔案不存在]")
            continue
        print(read_pdf(path) if fmt == "pdf" else read_docx(path))


if __name__ == "__main__":
    main()
